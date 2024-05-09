from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import pandas as pd
import sys

# Create a new Document
doc = Document()

case = sys.argv[1]

df = pd.read_csv(f"../results/{case}.csv")
df.rename(columns={'Unnamed: 0': 'item'}, inplace=True)
cleaned_df = df.dropna(subset=['score'])

# TO DROP
new_df = cleaned_df[cleaned_df["item"] != "NON-VERBAL FACILITATION SKILLS"]
new_df = new_df[new_df["item"] != "CLOSURE"]
sorted_df = new_df.sort_values(by='score', ascending=False)
# TO DROP

# sorted_df = cleaned_df.sort_values(by='score', ascending=False)
top_scores = sorted_df.head(4)
worst_score = sorted_df.tail(1)


# Set the title style and add the title to the document
title = doc.add_heading(level=0)
title_run = title.add_run('Narrative Feedback Form: Clinical Skills – Medical Interviewing')
title_run.font.size = Pt(14)
title_run.bold = True
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a horizontal line after the title
title._element.addnext(OxmlElement('w:p'))
# p = doc.add_paragraph()
# p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# run = p.add_run()
# run.add_break()
# p._element.addnext(OxmlElement('w:p'))


# Set up the table for Facilitator, Date, and Room entries
table = doc.add_table(rows=1, cols=3)
# table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Facilitator: GPT4'
hdr_cells[1].text = f'Date: {time.strftime("%m/%d/%Y")}'
hdr_cells[2].text = 'Room: Simulation Room 1'

# Add instructions paragraph
instructions_paragraph = doc.add_paragraph()
instructions_paragraph.add_run(
    "On this form, take note of the behavior-based reinforcing and modifying feedback for the "
    "student who performs an interview. Feedback will occur through self-reflection and peer and "
    "facilitator observation. At the end of the session, have the student take a photo of their feedback "
    "and leave this written copy in the room."
)
instructions_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add an image to the document and center it on the page
doc.add_picture('../results/frame.png', width=Inches(1.5), height=Inches(1.5))
last_paragraph = doc.paragraphs[-1]

# Center the image
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Create a table for the feedback form
feedback_table = doc.add_table(rows=1, cols=2)
feedback_table.style = 'Table Grid'

# Set column headers for reinforcing and modifying feedback
feedback_table.cell(0, 0).text = 'Reinforcing feedback'
feedback_table.cell(0, 1).text = 'Modifying feedback'


# Assuming the next row is where we add the top scores
# Insert each of the top scores as a new paragraph in the first column

# We will add a row for each top score
for index, row in top_scores.iterrows():
    cells = feedback_table.add_row().cells
    cells[0].text = f"{row[0]}:\n\n {row[2]}"
    cells[1].text = ''  # Modifying feedback is left blank

feedback_table.cell(1, 1).text = f"{worst_score['item'].values[0]}:\n\n {worst_score['explanation'].values[0]}"

# Set the column widths
for row in feedback_table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = Inches(3) if idx == 0 else Inches(3.25)

# # Add Student label
# student_label = doc.add_paragraph()
# student_label.add_run('Student 1:').bold = True
# student_label.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Add border to the student label
# student_label = student_label._element
pt = OxmlElement('w:pBdr')
# student_label.append(pt)
bottom_border = OxmlElement('w:bottom')
bottom_border.set(qn('w:val'), 'single')
pt.append(bottom_border)

# Ensure there's space after the table before next section
doc.add_paragraph()


# Save the document with improved formatting and boxes
file_path_improved = f'../results/narrative_feedback_form_{case}.docx'
doc.save(file_path_improved)


